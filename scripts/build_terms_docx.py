"""Generate properly-formatted bilingual Terms & Conditions Word documents.

- Arabic doc: full RTL, Cairo font (with Tahoma fallback), proper headings,
  bullet lists, numbered lists.
- English doc: LTR, Cairo font, same structure.
- Western (Latin) digits 0-9 are preserved as-is per project rules.
"""

from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.oxml import OxmlElement


AR_FONT = "Cairo"
AR_FONT_FALLBACK = "Tahoma"
EN_FONT = "Cairo"
EN_FONT_FALLBACK = "Calibri"


def set_paragraph_rtl(paragraph, rtl: bool) -> None:
    """Mark the paragraph as right-to-left at the XML level."""
    pPr = paragraph._p.get_or_add_pPr()
    bidi = pPr.find(qn("w:bidi"))
    if rtl:
        if bidi is None:
            bidi = OxmlElement("w:bidi")
            pPr.append(bidi)
        bidi.set(qn("w:val"), "1")
    elif bidi is not None:
        pPr.remove(bidi)


def set_run_font(run, *, primary: str, fallback: str, rtl: bool, size_pt: float, bold: bool = False, color: RGBColor | None = None) -> None:
    """Apply a font to a run for both ASCII and complex-script (Arabic) ranges."""
    run.font.name = primary
    run.font.size = Pt(size_pt)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), primary)
    rFonts.set(qn("w:hAnsi"), primary)
    rFonts.set(qn("w:cs"), primary)
    rFonts.set(qn("w:eastAsia"), fallback)
    if rtl:
        rtl_el = rPr.find(qn("w:rtl"))
        if rtl_el is None:
            rtl_el = OxmlElement("w:rtl")
            rPr.append(rtl_el)
        rtl_el.set(qn("w:val"), "1")
        # Mirror size for complex-script.
        szCs = rPr.find(qn("w:szCs"))
        if szCs is None:
            szCs = OxmlElement("w:szCs")
            rPr.append(szCs)
        szCs.set(qn("w:val"), str(int(size_pt * 2)))


def add_numbering_definitions(doc: Document) -> None:
    """Inject simple numbering definitions: numId=1 → bullets, numId=2 → decimal."""
    numbering_part = doc.part.numbering_part
    if numbering_part is None:
        from docx.oxml.numbering import CT_Numbering  # type: ignore
        numbering_part = doc.part.numbering_part  # creates it lazily in newer versions
    numbering_xml = numbering_part.element

    # Wipe existing dynamically-added abstractNums to keep things deterministic
    # but keep any document-default ones.
    existing_ids = {int(a.get(qn("w:abstractNumId"))) for a in numbering_xml.findall(qn("w:abstractNum"))}
    next_abs_id = max(existing_ids, default=-1) + 1

    bullet_abs_id = next_abs_id
    decimal_abs_id = next_abs_id + 1

    bullet_xml = f"""
<w:abstractNum xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:abstractNumId="{bullet_abs_id}">
  <w:multiLevelType w:val="hybridMultilevel"/>
  <w:lvl w:ilvl="0">
    <w:start w:val="1"/>
    <w:numFmt w:val="bullet"/>
    <w:lvlText w:val="•"/>
    <w:lvlJc w:val="left"/>
    <w:pPr><w:ind w:left="720" w:hanging="360"/></w:pPr>
    <w:rPr><w:rFonts w:ascii="Cairo" w:hAnsi="Cairo" w:cs="Cairo"/></w:rPr>
  </w:lvl>
</w:abstractNum>
"""
    decimal_xml = f"""
<w:abstractNum xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:abstractNumId="{decimal_abs_id}">
  <w:multiLevelType w:val="hybridMultilevel"/>
  <w:lvl w:ilvl="0">
    <w:start w:val="1"/>
    <w:numFmt w:val="decimal"/>
    <w:lvlText w:val="%1."/>
    <w:lvlJc w:val="left"/>
    <w:pPr><w:ind w:left="720" w:hanging="360"/></w:pPr>
  </w:lvl>
</w:abstractNum>
"""
    from lxml import etree
    bullet_el = etree.fromstring(bullet_xml)
    decimal_el = etree.fromstring(decimal_xml)

    # Insert abstractNums before any existing <w:num> entries.
    first_num = numbering_xml.find(qn("w:num"))
    if first_num is not None:
        first_num.addprevious(bullet_el)
        first_num.addprevious(decimal_el)
    else:
        numbering_xml.append(bullet_el)
        numbering_xml.append(decimal_el)

    # Pick concrete numIds that don't collide.
    existing_num_ids = {int(n.get(qn("w:numId"))) for n in numbering_xml.findall(qn("w:num"))}
    bullet_num_id = max(existing_num_ids, default=0) + 1
    decimal_num_id = bullet_num_id + 1

    for num_id, abs_id in ((bullet_num_id, bullet_abs_id), (decimal_num_id, decimal_abs_id)):
        num_xml = f"""
<w:num xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:numId="{num_id}">
  <w:abstractNumId w:val="{abs_id}"/>
</w:num>
"""
        numbering_xml.append(etree.fromstring(num_xml))

    return bullet_num_id, decimal_num_id


def set_list_item(paragraph, num_id: int, ilvl: int = 0) -> None:
    """Attach a paragraph to a numbering definition."""
    pPr = paragraph._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl_el = OxmlElement("w:ilvl")
    ilvl_el.set(qn("w:val"), str(ilvl))
    numId_el = OxmlElement("w:numId")
    numId_el.set(qn("w:val"), str(num_id))
    numPr.append(ilvl_el)
    numPr.append(numId_el)
    pPr.append(numPr)


@dataclass
class Block:
    kind: str  # 'h1' | 'subtitle' | 'h2' | 'p' | 'bullet' | 'num'
    text: str


SECTION_RE = re.compile(r"^(\d{1,2})\.\s*(.+)$")
NUM_ITEM_RE = re.compile(r"^(\d{1,2})\.(.+)$")


def parse_arabic_source(src: str) -> list[Block]:
    blocks: list[Block] = []
    in_numbered = False
    for raw in src.splitlines():
        line = raw.strip()
        if not line:
            in_numbered = False
            continue
        if line == "الشروط والأحكام":
            blocks.append(Block("h1", line))
            continue
        if line.startswith("شركة العربة الفاخرة") or line.startswith("منصة"):
            blocks.append(Block("subtitle", line.rstrip("—").strip()))
            continue
        if line.startswith("تاريخ السريان"):
            blocks.append(Block("subtitle", line))
            continue
        if line.startswith("•"):
            blocks.append(Block("bullet", line.lstrip("•").strip().lstrip('"').rstrip('"') if False else line.lstrip("•").strip()))
            in_numbered = False
            continue
        # Numbered list item INSIDE section 5 (toplevel sections like "1." 2." are
        # treated as section headings; ordered items appear AFTER the "توافق على
        # عدم القيام بما يلي:" cue paragraph and before section 6).
        m_num = NUM_ITEM_RE.match(line)
        if m_num and in_numbered:
            blocks.append(Block("num", m_num.group(2).strip()))
            continue
        m_section = SECTION_RE.match(line)
        if m_section:
            num = int(m_section.group(1))
            title = m_section.group(2).strip()
            blocks.append(Block("h2", f"{num}. {title}"))
            in_numbered = False
            continue
        # Cue line that introduces the inline numbered list in section 5.
        if line.endswith("بما يلي:"):
            blocks.append(Block("p", line))
            in_numbered = True
            continue
        blocks.append(Block("p", line))
    return blocks


def parse_english_source(src: str) -> list[Block]:
    blocks: list[Block] = []
    for raw in src.splitlines():
        line = raw.rstrip()
        if not line.strip():
            continue
        if line.startswith("# "):
            blocks.append(Block("h1", line[2:].strip()))
            continue
        if line.startswith("## "):
            blocks.append(Block("h2", line[3:].strip()))
            continue
        if line.startswith("- "):
            blocks.append(Block("bullet", line[2:].strip()))
            continue
        m = re.match(r"^(\d+)\.\s+(.+)$", line)
        if m:
            blocks.append(Block("num", m.group(2).strip()))
            continue
        # Bold subtitle line like **Luxury Cart Company — WORKFORCE Platform**
        if line.startswith("**") and line.endswith("**"):
            blocks.append(Block("subtitle", line.strip("*")))
            continue
        if line.startswith("Effective date"):
            blocks.append(Block("subtitle", line))
            continue
        blocks.append(Block("p", line))
    return blocks


def strip_md_bold(text: str) -> tuple[list[tuple[str, bool]], ]:
    """Return list of (segment, is_bold) for inline **bold** markdown."""
    parts: list[tuple[str, bool]] = []
    last = 0
    for m in re.finditer(r"\*\*([^*]+)\*\*", text):
        if m.start() > last:
            parts.append((text[last:m.start()], False))
        parts.append((m.group(1), True))
        last = m.end()
    if last < len(text):
        parts.append((text[last:], False))
    if not parts:
        parts.append((text, False))
    return parts


def render_doc(blocks: list[Block], out_path: Path, *, rtl: bool, primary_font: str, fallback_font: str) -> None:
    doc = Document()

    # Page setup: A4, generous margins.
    for section in doc.sections:
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)
        # Mark the section as RTL so headers/footers flow correctly.
        sectPr = section._sectPr
        bidi = sectPr.find(qn("w:bidi"))
        if rtl and bidi is None:
            bidi = OxmlElement("w:bidi")
            sectPr.append(bidi)

    bullet_num_id, decimal_num_id = add_numbering_definitions(doc)

    align_start = WD_ALIGN_PARAGRAPH.RIGHT if rtl else WD_ALIGN_PARAGRAPH.LEFT

    def add_paragraph(text: str, *, size: float, bold: bool = False, alignment=None, color: RGBColor | None = None, list_num_id: int | None = None, space_before: float = 0, space_after: float = 6):
        p = doc.add_paragraph()
        p.alignment = alignment if alignment is not None else align_start
        set_paragraph_rtl(p, rtl)
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.4
        if list_num_id is not None:
            set_list_item(p, list_num_id)
        for seg, seg_bold in strip_md_bold(text):
            run = p.add_run(seg)
            set_run_font(run, primary=primary_font, fallback=fallback_font, rtl=rtl, size_pt=size, bold=bold or seg_bold, color=color)
        return p

    for blk in blocks:
        if blk.kind == "h1":
            add_paragraph(
                blk.text, size=22, bold=True,
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                color=RGBColor(0x11, 0x18, 0x27),
                space_before=0, space_after=10,
            )
        elif blk.kind == "subtitle":
            add_paragraph(
                blk.text, size=12, bold=True,
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                color=RGBColor(0x4B, 0x55, 0x63),
                space_after=4,
            )
        elif blk.kind == "h2":
            add_paragraph(
                blk.text, size=14, bold=True,
                color=RGBColor(0x16, 0x36, 0x6F),
                space_before=14, space_after=4,
            )
        elif blk.kind == "bullet":
            add_paragraph(blk.text, size=11, list_num_id=bullet_num_id, space_after=2)
        elif blk.kind == "num":
            add_paragraph(blk.text, size=11, list_num_id=decimal_num_id, space_after=2)
        else:
            add_paragraph(blk.text, size=11)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    print(f"Wrote {out_path}")


def main() -> None:
    root = Path(__file__).resolve().parent.parent
    ar_md = (root / "attached_assets" / "terms_conditions_ar_2026-04-30.md").read_text(encoding="utf-8")
    en_md = (root / "attached_assets" / "terms_conditions_en_2026-04-30.md").read_text(encoding="utf-8")

    ar_blocks = parse_arabic_source(ar_md)
    en_blocks = parse_english_source(en_md)

    render_doc(
        ar_blocks,
        out_path=root / "attached_assets" / "terms_conditions_ar_2026-04-30.docx",
        rtl=True,
        primary_font=AR_FONT,
        fallback_font=AR_FONT_FALLBACK,
    )
    render_doc(
        en_blocks,
        out_path=root / "attached_assets" / "terms_conditions_en_2026-04-30.docx",
        rtl=False,
        primary_font=EN_FONT,
        fallback_font=EN_FONT_FALLBACK,
    )


if __name__ == "__main__":
    main()
